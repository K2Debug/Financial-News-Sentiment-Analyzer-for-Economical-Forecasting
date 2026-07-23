"""
Generate EF-02 Comprehensive Academic Report as Word document.
"""
from __future__ import annotations

import os
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
DATA = ROOT.parent / "research" / "data" / "processed"
FIG_DIR = ROOT / "figures"
OUTPUT = ROOT / "EF02_Project_Report.docx"

FONT = "Times New Roman"
BODY_SIZE = 12


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size in [(1, 14), (2, 13), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        if level == 1:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_para(doc: Document, text: str, bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(BODY_SIZE)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None) -> None:
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = cap.add_run(caption)
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(BODY_SIZE)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = FONT
                run.font.size = Pt(BODY_SIZE)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.name = FONT
                    run.font.size = Pt(BODY_SIZE)
    doc.add_paragraph()


def add_toc(doc: Document) -> None:
    add_heading(doc, "TABLE OF CONTENTS", 1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and select Update Field in Microsoft Word to refresh the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)
    doc.add_page_break()

    add_heading(doc, "LIST OF TABLES", 1)
    tables_list = [
        "Table 1.1: Value Proposition Canvas for EF-02",
        "Table 3.1: Methodology Table",
        "Table 5.1: Model Benchmarking Results",
        "Table 5.2: Functional Testing Results",
        "Table 5.3: Pearson Correlation Results",
    ]
    for t in tables_list:
        add_para(doc, t)
    doc.add_page_break()

    add_heading(doc, "LIST OF FIGURES", 1)
    figures_list = [
        "Figure 5.1: Headline relevance split",
        "Figure 5.2: Monthly sentiment distribution",
        "Figure 5.3: Positive and Negative sentiment versus USD/TZS exchange rate",
        "Figure 5.4: Pearson correlation matrix",
    ]
    for f in figures_list:
        add_para(doc, f)
    doc.add_page_break()


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.5) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = FONT
    r.font.size = Pt(10)
    doc.add_paragraph()


def export_figures() -> dict[str, Path]:
    FIG_DIR.mkdir(exist_ok=True)
    labelled = pd.read_csv(DATA / "tz_headlines_labelled.csv")
    viz = pd.read_csv(DATA / "Visualization_Data.csv")
    viz["Net_Sentiment"] = viz["Positive %"] - viz["Negative %"]
    paths: dict[str, Path] = {}

    # Figure 5.1: Relevance donut
    fig, ax = plt.subplots(figsize=(6, 4))
    rel = labelled["relevant"].value_counts()
    labels = ["Relevant", "Irrelevant"]
    sizes = [rel.get(True, 0), rel.get(False, 0)]
    pcts = [s / sum(sizes) * 100 for s in sizes]
    ax.pie(
        sizes,
        labels=[f"{l}\n({p:.1f}%)" for l, p in zip(labels, pcts)],
        colors=["#2196F3", "#E0E0E0"],
        startangle=90,
        wedgeprops=dict(width=0.55),
    )
    ax.set_title("Headline Relevance Split (n=5,310)")
    p1 = FIG_DIR / "fig5_1_relevance.png"
    fig.savefig(p1, dpi=150, bbox_inches="tight")
    plt.close(fig)
    paths["relevance"] = p1

    # Figure 5.2: Monthly sentiment stacked bar
    fig, ax = plt.subplots(figsize=(10, 4))
    x = range(len(viz))
    ax.bar(x, viz["Neutral %"], label="Neutral", color="#9E9E9E", width=0.8)
    ax.bar(x, viz["Positive %"], bottom=viz["Neutral %"], label="Positive", color="#2196F3", width=0.8)
    ax.bar(
        x,
        viz["Negative %"],
        bottom=viz["Neutral %"] + viz["Positive %"],
        label="Negative",
        color="#F44336",
        width=0.8,
    )
    ax.set_xticks(x[::3])
    ax.set_xticklabels(viz["YearMonth"].iloc[::3], rotation=45, ha="right")
    ax.set_ylabel("Share (%)")
    ax.set_title("Monthly Sentiment Distribution (2022 to 2024)")
    ax.legend()
    fig.tight_layout()
    p2 = FIG_DIR / "fig5_2_sentiment.png"
    fig.savefig(p2, dpi=150, bbox_inches="tight")
    plt.close(fig)
    paths["sentiment"] = p2

    # Figure 5.3: Dual axis Positive vs USD/TZS
    fig, axes = plt.subplots(1, 2, figsize=(12, 4))
    x = range(len(viz))
    for ax, col, color, title in [
        (axes[0], "Positive %", "#2196F3", "Positive % vs USD/TZS Rate"),
        (axes[1], "Negative %", "#F44336", "Negative % vs USD/TZS Rate"),
    ]:
        ax2 = ax.twinx()
        ax.plot(x, viz[col], color=color, marker="o", markersize=3, label=col)
        ax2.plot(x, viz["USD/TZS_Rate"], color="gray", linestyle="--", label="USD/TZS")
        ax.set_xticks(x[::4])
        ax.set_xticklabels(viz["YearMonth"].iloc[::4], rotation=45, ha="right", fontsize=8)
        ax.set_title(title)
        ax.set_ylabel("Sentiment %")
        ax2.set_ylabel("TZS per USD")
    fig.tight_layout()
    p3 = FIG_DIR / "fig5_3_dual_axis.png"
    fig.savefig(p3, dpi=150, bbox_inches="tight")
    plt.close(fig)
    paths["dual_axis"] = p3

    # Figure 5.4: Correlation heatmap
    corr_cols = [
        "Positive %",
        "Negative %",
        "Neutral %",
        "Net_Sentiment",
        "Inflation %",
        "USD/TZS_Rate",
        "Rate_Change_%",
    ]
    corr = viz[corr_cols].corr(method="pearson")
    fig, ax = plt.subplots(figsize=(8, 6))
    sns.heatmap(corr, annot=True, fmt=".2f", cmap="RdBu_r", center=0, ax=ax, linewidths=0.5)
    ax.set_title("Pearson Correlation Matrix")
    fig.tight_layout()
    p4 = FIG_DIR / "fig5_4_correlation.png"
    fig.savefig(p4, dpi=150, bbox_inches="tight")
    plt.close(fig)
    paths["correlation"] = p4

    return paths


def build_cover(doc: Document) -> None:
    lines = [
        ("ARDHI UNIVERSITY", True, True),
        (
            "SCHOOL OF EARTH SCIENCES, REAL ESTATE, BUSINESS STUDIES AND INFORMATICS (SERBI)",
            False,
            True,
        ),
        ("DEPARTMENT OF COMPUTER SYSTEMS AND MATHEMATICS (CSM)", False, True),
        ("", False, True),
        ("PROGRAMME: DATA SCIENCE AND ARTIFICIAL INTELLIGENCE", False, True),
        ("COURSE CODE: DA 181", False, True),
        (
            "COURSE NAME: PROJECT 1: PROBLEM SOLVING USING FUNDAMENTAL CONCEPTS OF DATA SCIENCE AND ARTIFICIAL INTELLIGENCE",
            False,
            True,
        ),
        ("LEVEL: UQF 8 / FIRST YEAR", False, True),
        ("INSTRUCTORS: DR. AHMAD AND MR. KAMALI", False, True),
        ("SUBMISSION DATE: 19TH FEBRUARY, 2026", False, True),
        ("", False, True),
        (
            "PROJECT TITLE: EF-02: FINANCIAL NEWS SENTIMENT ANALYSER FOR ECONOMIC FORECASTING",
            True,
            True,
        ),
        ("", False, True),
        ("GROUP MEMBERS", True, True),
    ]
    for text, bold, center in lines:
        if text:
            add_para(doc, text, bold=bold, center=center)

    add_table(
        doc,
        ["S/N", "FULL NAMES", "REG. NO."],
        [
            ["1", "KITULA, KITULA-MASAGANYA ABDUL", "35552/T.2025"],
            ["2", "ILEKA, REVOCATUS OREST", "37058/T.2025"],
            ["3", "LYELLU, SABRINA RASHID", "37463/T.2025"],
            ["4", "IBRAHIM, GHADHAL NKRUMAH", "37254/T.2025"],
        ],
    )
    doc.add_page_break()


def build_front_matter(doc: Document) -> None:
    add_heading(doc, "DECLARATION", 1)
    add_para(
        doc,
        "We, the members of Group 8, hereby declare that this report is our own work and effort. "
        "The work in this report was carried out in accordance with the Regulations of Ardhi University. "
        "We have faithfully acknowledged, given credit to, and referred to the research workers wherever "
        "their work has been cited in the text and the body of this project. We further certify that we "
        "have not wilfully lifted another person's paragraph, text, data, or results reported in journals, "
        "books, magazines, reports, dissertations, theses, or available at websites and included them in "
        "this project and cited them as our own work.",
    )
    add_table(
        doc,
        ["S/N", "STUDENT'S NAME", "SIGNATURE", "DATE"],
        [
            ["1", "KITULA, KITULA-MASAGANYA ABDUL", "", ""],
            ["2", "ILEKA, REVOCATUS OREST", "", ""],
            ["3", "LYELLU, SABRINA RASHID", "", ""],
            ["4", "IBRAHIM, GHADHAL NKRUMAH", "", ""],
        ],
    )
    doc.add_page_break()

    add_heading(doc, "CERTIFICATION", 1)
    add_para(
        doc,
        "The undersigned certify that they have read and hereby recommend for acceptance by Ardhi University "
        "a project titled: EF-02: Financial News Sentiment Analyser for Economic Forecasting, in fulfilment "
        "of the requirements for the accomplishment of first year studies at Ardhi University.",
    )
    add_para(doc, "___________________________", center=False)
    add_para(doc, "Dr. Ahmad (Supervisor)")
    add_para(doc, "Date: _______________")
    add_para(doc, "")
    add_para(doc, "___________________________", center=False)
    add_para(doc, "Mr. Kamali (Supervisor)")
    add_para(doc, "Date: _______________")
    doc.add_page_break()

    add_heading(doc, "ACKNOWLEDGEMENT", 1)
    add_para(
        doc,
        "First and foremost, we thank Almighty God for giving us the strength and guidance to complete this project. "
        "We express our sincere gratitude to our supervisors, Dr. Ahmad and Mr. Kamali, for their advice, feedback, "
        "and support throughout the development of EF-02. Their guidance helped us structure our pipeline and present "
        "our findings in a clear academic manner.",
    )
    add_para(
        doc,
        "We also thank Ardhi University and the Department of Computer Systems and Mathematics for providing the "
        "learning environment and resources that made this work possible. We acknowledge Daily News Tanzania and "
        "The Citizen for publishing the financial news used in this study, the National Bureau of Statistics (NBS) "
        "for CPI data, and Investing.com for historical USD/TZS exchange rate data. Finally, we thank Groq for "
        "providing API access to large language models used in headline classification.",
    )
    doc.add_page_break()

    add_heading(doc, "LIST OF ABBREVIATIONS AND ACRONYMS", 1)
    abbrevs = [
        ("ARU", "Ardhi University"),
        ("BOT", "Bank of Tanzania"),
        ("CPI", "Consumer Price Index"),
        ("CRISP-DM", "Cross Industry Standard Process for Data Mining"),
        ("DSE", "Dar es Salaam Stock Exchange"),
        ("FX", "Foreign Exchange"),
        ("LLM", "Large Language Model"),
        ("NBS", "National Bureau of Statistics"),
        ("NLP", "Natural Language Processing"),
        ("TZS", "Tanzanian Shilling"),
        ("USD", "United States Dollar"),
        ("VPC", "Value Proposition Canvas"),
    ]
    for short, full in abbrevs:
        add_para(doc, f"{short} : {full}")
    doc.add_page_break()

    add_heading(doc, "ABSTRACT", 1)
    abstract_paras = [
        "This project investigates whether the sentiment of Tanzanian financial news headlines is related to "
        "macroeconomic movements, specifically the USD/TZS exchange rate and headline inflation, between 2022 and 2024. "
        "Financial media often reflects and shapes public expectations about the economy. However, few studies have "
        "applied modern Natural Language Processing methods to Tanzanian financial news using local economic data.",
        "A six stage data science pipeline was built: web scraping from Daily News and The Citizen, data cleaning, "
        "multi model benchmarking, LLM based classification via the Groq API, monthly data consolidation with NBS "
        "CPI and USD/TZS series, and visualisation with Pearson correlation analysis. A total of 5,310 headlines were "
        "collected. After relevance filtering, 72.2% (3,836 headlines) were classified into nine economic categories "
        "and three sentiment labels based on implied economic outcome.",
        "Monthly aggregates were produced for 33 months (March 2022 to December 2024). January and February 2022 were "
        "missing from the exchange rate source. Model benchmarking showed that rule based tools (TextBlob at 28.2% and "
        "VADER at 50.0%) and FinBERT at 66.7% performed poorly on Tanzanian financial text, while large language models "
        "achieved about 93% accuracy on a labelled sample.",
        "Correlation results showed that sentiment had its strongest statistical link with the USD/TZS exchange rate "
        "level. Net Sentiment (Positive share minus Negative share) correlated with USD/TZS at r = 0.583 (p = 0.0004). "
        "Negative sentiment correlated at r = 0.484 (p = 0.0043) and Positive sentiment at r = 0.438 (p = 0.0108). "
        "Net Sentiment also showed a moderate inverse link with inflation at r = 0.413 (p = 0.017). Month on month "
        "rate changes did not show significant correlation with sentiment.",
        "The project demonstrates that Tanzanian financial news sentiment can serve as a concurrent indicator of "
        "exchange rate movement, though limitations include neutral dominance in labels, a modest monthly sample size, "
        "and the fact that correlation does not prove causation. The pipeline offers a replicable framework for future "
        "African financial NLP research.",
    ]
    for p in abstract_paras:
        add_para(doc, p)
    doc.add_page_break()
    add_toc(doc)


def build_chapter1(doc: Document) -> None:
    add_heading(doc, "CHAPTER ONE: INTRODUCTION", 1)

    add_heading(doc, "1.1 General Introduction", 2)
    add_para(
        doc,
        "Sentiment analysis, also called opinion mining, is a branch of Natural Language Processing (NLP) that "
        "identifies subjective information in text. In finance, researchers study whether the tone of news, reports, "
        "and social media can help explain or predict economic behaviour. Loughran and McDonald (2011) showed that "
        "general purpose sentiment dictionaries perform poorly on financial text because many words carry special "
        "meanings in economics. Baker, Bloom, and Davis (2016) built an Economic Policy Uncertainty index partly from "
        "newspaper tone, which is now used in many countries.",
    )
    add_para(
        doc,
        "In Tanzania, English language outlets such as Daily News and The Citizen publish regular business and financial "
        "news covering monetary policy, banking, trade, agriculture, energy, and foreign exchange. Despite this rich "
        "media output, computational analysis of Tanzanian financial news remains limited. Adelani et al. (2021) noted "
        "that African language and domain NLP resources are underdeveloped. This project applies data science methods "
        "to Tanzanian financial headlines from 2022 to 2024 and tests whether monthly sentiment patterns relate to "
        "official macroeconomic indicators published by the Bank of Tanzania and the National Bureau of Statistics.",
    )

    add_heading(doc, "1.2 Statement of the Problem", 2)
    add_para(
        doc,
        "Financial news shapes how the public, investors, and policymakers understand the economy. Headlines about "
        "currency weakness, inflation, or banking performance can influence expectations even before official statistics "
        "are released. In Tanzania, this information exists in large volumes online, but it is mostly unstructured text.",
    )
    add_para(
        doc,
        "Existing sentiment research focuses on Western markets and tools trained on Western financial language. "
        "Rule based models and even FinBERT struggle with Tanzanian economic phrasing. There is also no published study "
        "that correlates Tanzanian financial media sentiment with BOT reported inflation and the USD/TZS exchange rate "
        "over a multi year period.",
    )
    add_para(
        doc,
        "This project addresses the gap by building a reproducible pipeline that scrapes local news, classifies headlines "
        "using a large language model, merges results with official economic data, and tests statistical relationships "
        "through correlation analysis and visualisation.",
    )

    add_heading(doc, "1.3 Objectives", 2)
    add_heading(doc, "1.3.1 General Objective", 3)
    add_para(
        doc,
        "The general objective of this project is to investigate whether sentiment expressed in Tanzanian financial "
        "news media is meaningfully correlated with movements in the USD/TZS exchange rate and CPI inflation over "
        "the period 2022 to 2024.",
    )
    add_heading(doc, "1.3.2 Specific Objectives", 3)
    objectives = [
        "To apply a large language model via the Groq API to classify 5,310 Tanzanian financial news headlines by topic category and economic sentiment label.",
        "To compute monthly average sentiment scores aggregated from the labelled headline dataset across the study period.",
        "To correlate monthly sentiment scores with NBS published CPI inflation and USD/TZS exchange rate data using Pearson correlation.",
        "To visualise the relationship between media sentiment trends and macroeconomic movements through time series plots and correlation charts.",
        "To evaluate the performance and suitability of NLP approaches, including rule based tools, FinBERT, and LLMs, for Tanzanian financial text.",
    ]
    for i, obj in enumerate(objectives, 1):
        add_para(doc, f"{i}. {obj}")

    add_heading(doc, "1.3.3 Value Proposition Canvas", 3)
    add_para(
        doc,
        "A Value Proposition Canvas (VPC) is a tool that links what users need with what a product offers. It helps "
        "clarify who benefits from a system and why it matters. For EF-02, the canvas below maps stakeholders, their "
        "problems, desired outcomes, and the project deliverables.",
    )
    add_table(
        doc,
        ["VPC Element", "Description for EF-02"],
        [
            ["Customer segments", "Policymakers, economists, investors, researchers, and data science students"],
            ["Customer jobs", "Track currency pressure, inflation mood, and sector trends from media before official data"],
            ["Pains", "Delayed official statistics, Western biased NLP tools, unstructured news data"],
            ["Gains", "Monthly sentiment index, category breakdown, correlation evidence with macro indicators"],
            ["Products and services", "Six notebook pipeline producing labelled headlines and Visualization_Data.csv"],
            ["Pain relievers", "Local news sources, LLM classification tuned for economic outcome, monthly aggregation"],
            ["Gain creators", "Charts, correlation tables, and a replicable framework for East African financial NLP"],
        ],
        caption="Table 1.1: Value Proposition Canvas for EF-02",
    )

    add_heading(doc, "1.4 Research Questions", 2)
    questions = [
        "Can a large language model accurately classify Tanzanian financial headlines by topic and economic sentiment?",
        "Do monthly sentiment scores show a statistical relationship with the USD/TZS exchange rate?",
        "Do monthly sentiment scores show a statistical relationship with CPI inflation?",
        "Which NLP approach performs best on Tanzanian financial text compared to rule based and transformer models?",
        "What limitations affect the reliability of the sentiment signal in this study?",
    ]
    for i, q in enumerate(questions, 1):
        add_para(doc, f"RQ{i}: {q}")

    add_heading(doc, "1.5 Significance of the Study", 2)
    sig = [
        "It contributes local evidence on whether financial media sentiment relates to Tanzanian macro indicators.",
        "It demonstrates a practical data science pipeline that students and researchers can extend to other sectors.",
        "It highlights the need for African context NLP tools rather than relying only on Western trained models.",
        "It supports early awareness of currency and inflation trends for academic and policy discussion.",
        "It builds skills in scraping, prompt engineering, time series analysis, and visualisation.",
    ]
    add_para(doc, "Upon completing this study, the following benefits are expected:")
    for s in sig:
        add_para(doc, f"• {s}")

    add_heading(doc, "1.6 Structure of the Report", 2)
    add_para(
        doc,
        "Chapter One introduces the problem, objectives, research questions, and significance. Chapter Two reviews "
        "related literature and identifies the research gap. Chapter Three explains the methodology and tools used. "
        "Chapter Four presents system requirements and architecture. Chapter Five describes implementation, results, "
        "and testing. Chapter Six provides conclusions, limitations, and recommendations.",
    )
    doc.add_page_break()


def build_chapter2(doc: Document) -> None:
    add_heading(doc, "CHAPTER TWO: LITERATURE REVIEW", 1)

    add_heading(doc, "2.1 Introduction", 2)
    add_para(
        doc,
        "A literature review surveys existing research to show what is already known about a topic and where gaps remain. "
        "It helps justify the methods chosen and positions this project within academic discourse. This chapter reviews "
        "studies on financial sentiment analysis, NLP methods, links between media tone and macroeconomic indicators, "
        "and research in the African context.",
    )

    add_heading(doc, "2.2 Related Studies", 2)

    add_heading(doc, "2.2.1 Sentiment Analysis in Financial Contexts", 3)
    add_para(
        doc,
        "Loughran and McDonald (2011) developed a finance specific lexicon and showed that general dictionaries misread "
        "many financial terms. Tetlock (2007) found that the tone of Wall Street Journal columns had a predictive "
        "relationship with stock market returns, suggesting that media pessimism can move prices. Bollen, Mao, and Zeng "
        "(2011) demonstrated that public mood on Twitter could predict movements in the Dow Jones Industrial Average. "
        "These studies established media sentiment as a legitimate variable in economic modelling.",
    )

    add_heading(doc, "2.2.2 NLP Approaches: From Rule Based to Large Language Models", 3)
    add_para(
        doc,
        "Early sentiment systems used lexicon and rule based tools such as VADER (Hutto and Gilbert, 2014) and TextBlob. "
        "These methods are fast and interpretable but struggle with context, negation, and domain specific language. "
        "For example, the phrase 'the Bank of Tanzania holds rates steady' may signal stability (positive economic outcome) "
        "or lack of action (neutral), which simple word lists cannot reliably distinguish.",
    )
    add_para(
        doc,
        "Transformer models changed the field. Devlin et al. (2018) introduced BERT, a context aware pre trained model. "
        "Araci (2019) fine tuned FinBERT on financial text and improved performance on financial sentiment benchmarks. "
        "More recently, instruction following LLMs such as Llama 3, accessed through API providers like Groq, allow "
        "researchers to classify relevance, category, and sentiment in a single step, reducing pipeline complexity.",
    )

    add_heading(doc, "2.2.3 Sentiment and Macroeconomic Indicators", 3)
    add_para(
        doc,
        "Shapiro, Sudhof, and Wilson (2022) built a monthly news sentiment index from thousands of newspaper articles "
        "and found it predicted GDP growth and unemployment in the United States. Baker et al. (2016) showed that policy "
        "uncertainty derived from media tone affects investment and employment. These frameworks inspire the present "
        "study, which tests similar ideas using Tanzanian news and local CPI and exchange rate data.",
    )

    add_heading(doc, "2.2.4 Financial News in Sub-Saharan Africa", 3)
    add_para(
        doc,
        "Digital media in Sub-Saharan Africa has grown rapidly. Outlets such as The Citizen and Daily News produce steady "
        "streams of financial reporting. Muriungi and Kimani (2020) examined Kenyan financial news sentiment and found "
        "a moderate correlation with the KES/USD exchange rate. However, direct studies for Tanzania using TZS, BOT policy "
        "data, and local media sources remain scarce. Adelani et al. (2021) emphasised that African NLP resources are "
        "critically underdeveloped, which supports the value of projects that evaluate tools in African financial settings.",
    )

    add_heading(doc, "2.3 Research Gap", 2)
    gaps = [
        "Most sentiment and economy studies focus on Western or East Asian markets, not Tanzania.",
        "Prior work often treats topic classification and sentiment scoring as separate steps; single pass LLM pipelines are rarely evaluated on African financial news.",
        "Tanzanian media sentiment has not been systematically correlated with NBS inflation and USD/TZS data.",
        "Public datasets for financial sentiment are dominated by Reuters, Bloomberg, and social media, not East African outlets.",
    ]
    add_para(doc, "The literature review reveals four main gaps that EF-02 addresses:")
    for i, g in enumerate(gaps, 1):
        add_para(doc, f"{i}. {g}")
    add_para(
        doc,
        "This project fills these gaps by scraping Tanzanian news, classifying headlines with a Groq hosted LLM, merging "
        "with official macro data, and reporting correlation results for 2022 to 2024.",
    )
    doc.add_page_break()


def build_chapter3(doc: Document) -> None:
    add_heading(doc, "CHAPTER THREE: METHODOLOGY", 1)

    add_heading(doc, "3.1 Introduction", 2)
    add_para(
        doc,
        "This chapter describes how the project was carried out. It explains the overall methodology, how requirements "
        "were gathered, how the system was designed and implemented, and how results were tested and validated.",
    )

    add_heading(doc, "3.2 Selected Methodology", 2)

    add_heading(doc, "3.2.1 General Methodology", 3)
    add_para(
        doc,
        "The project followed CRISP-DM (Cross Industry Standard Process for Data Mining), adapted with an iterative "
        "rebuild approach. CRISP-DM covers business understanding, data understanding, data preparation, modelling, "
        "evaluation, and deployment. This lifecycle fits a pipeline that moves from raw news text to labelled data, "
        "merged monthly series, and visual insights. The team first benchmarked models on sample data, then scraped "
        "full headlines, then ran production classification, as recorded in the project decision log.",
    )

    add_heading(doc, "3.2.2 Methodology for Gathering User Requirements", 3)
    add_para(
        doc,
        "Requirements were gathered through literature review, inspection of Tanzanian news websites, and comparison "
        "of data availability. Daily News and The Citizen were selected because they cover 2022 to 2024 and allow "
        "requests based scraping. IPP Media was excluded because its archive lacked full range coverage and its category "
        "filters could not be accessed without a headless browser.",
    )

    add_heading(doc, "3.2.3 Methodology for System Design", 3)
    add_para(
        doc,
        "The system was designed as a numbered notebook pipeline with clear inputs and outputs per stage. Architecture "
        "documentation mapped each notebook to one step: scraper, cleaning, benchmarking, classifier, consolidation, "
        "and visualisation. Data files were stored in raw and processed folders to keep source data separate from derived results.",
    )

    add_heading(doc, "3.2.4 Methodology for System Implementation", 3)
    add_para(
        doc,
        "Implementation used Python 3 in Jupyter notebooks. Web scraping relied on requests and BeautifulSoup. Data "
        "processing used pandas. The classifier used the OpenAI SDK pointed at Groq's API endpoint with Llama models. "
        "Charts were produced with matplotlib and seaborn. API keys were stored in a .env file and loaded with python-dotenv.",
    )

    add_heading(doc, "3.2.5 Methodology for System Testing and Validation", 3)
    add_para(
        doc,
        "Testing included multi model benchmarking against a labelled sample, batch validation of LLM JSON output "
        "(correct row count and valid labels), functional checks on each pipeline stage, and Pearson correlation with "
        "p values to test statistical significance of sentiment versus economic variable pairs.",
    )

    add_heading(doc, "3.3 Methodology Table", 2)
    add_table(
        doc,
        ["Specific Objective", "Methodology", "Tools", "Deliverable"],
        [
            [
                "Gather financial news headlines",
                "Web scraping with checkpointing every 30 pages",
                "requests, BeautifulSoup",
                "dailynews_raw.csv, citizen_raw.csv",
            ],
            [
                "Clean and unify headline data",
                "Merge, deduplicate, normalise dates",
                "pandas",
                "tz_headlines_clean.csv",
            ],
            [
                "Select best classification approach",
                "Multi model accuracy comparison",
                "TextBlob, VADER, FinBERT, Groq LLMs",
                "benchmarking_results.csv",
            ],
            [
                "Classify headlines by category and sentiment",
                "Single pass 3 in 1 LLM prompt per batch",
                "Groq API, Llama models",
                "tz_headlines_labelled.csv",
            ],
            [
                "Merge news with macroeconomic data",
                "Monthly aggregation and join on YearMonth",
                "pandas",
                "Visualization_Data.csv",
            ],
            [
                "Analyse sentiment vs economy",
                "Time series plots and Pearson correlation",
                "matplotlib, seaborn, scipy",
                "06_visualisation.ipynb outputs",
            ],
        ],
        caption="Table 3.1: Methodology Table",
    )
    doc.add_page_break()


def build_chapter4(doc: Document) -> None:
    add_heading(doc, "CHAPTER FOUR: SYSTEM ANALYSIS AND DESIGN", 1)

    add_heading(doc, "4.1 Introduction", 2)
    add_para(
        doc,
        "This chapter defines what the system needs in terms of data, software, hardware, and user expectations. "
        "It also presents the architecture that connects each pipeline stage from raw news to final analysis.",
    )

    add_heading(doc, "4.2 Requirement Analysis", 2)
    add_para(
        doc,
        "Requirement analysis identifies what a system must do and what conditions it must meet. This section covers "
        "system requirements (data, software, hardware) and user requirements (functional and non functional).",
    )

    add_heading(doc, "4.2.1 System Requirements", 3)
    add_para(
        doc,
        "System requirements describe the resources and environment needed to run the pipeline. They include data, "
        "software, and hardware components.",
    )

    add_heading(doc, "4.2.1.1 Data Requirements", 4)
    add_para(doc, "The following data types were collected and used:")
    data_reqs = [
        ("News headlines", "Scraped from Daily News and The Citizen business sections (2022 to 2024). Fields: date, headline, url, source."),
        ("USD/TZS exchange rate", "Daily data from Investing.com, resampled to monthly averages with month on month change."),
        ("CPI inflation", "Monthly headline inflation from NBS Tanzania CPI summary (2022 to 2024)."),
        ("Labelled sample", "Synthetic and test headlines used for model benchmarking before full classification."),
    ]
    for name, desc in data_reqs:
        add_para(doc, f"{name}: {desc}")

    add_heading(doc, "4.2.1.2 Software Requirements", 4)
    sw = [
        "Python 3: main programming language for all notebooks.",
        "Jupyter Notebook: interactive development and reproducible execution.",
        "pandas: data cleaning, merging, and monthly aggregation.",
        "requests and BeautifulSoup: HTTP requests and HTML parsing for scraping.",
        "OpenAI SDK (Groq endpoint): LLM API calls for classification.",
        "matplotlib and seaborn: charts and correlation heatmap.",
        "python-dotenv: secure loading of API keys from .env file.",
        "scipy: Pearson correlation and p value calculation.",
    ]
    for s in sw:
        add_para(doc, f"• {s}")

    add_heading(doc, "4.2.1.3 Hardware Requirements", 4)
    add_para(
        doc,
        "A standard laptop or desktop computer with at least 8 GB RAM, internet connection for web scraping and API "
        "calls, and sufficient storage for CSV datasets (under 50 MB processed) was sufficient. No specialised GPU was "
        "required because LLM inference ran on Groq cloud servers.",
    )

    add_heading(doc, "4.2.2 User Requirements", 3)
    add_para(
        doc,
        "User requirements describe what stakeholders expect the system to achieve. They are grouped into functional "
        "requirements (what the system does) and non functional requirements (how well it performs).",
    )

    add_heading(doc, "4.2.2.1 Functional Requirements", 4)
    func = [
        "Scrape and store headlines with date, source, and article URL.",
        "Filter irrelevant headlines and assign economic category and sentiment.",
        "Aggregate relevant headlines to monthly sentiment percentages.",
        "Merge sentiment data with CPI inflation and USD/TZS rate on a common YearMonth key.",
        "Produce charts and correlation statistics for interpretation.",
        "Support resume and checkpoint during long scraping runs.",
    ]
    for i, f in enumerate(func, 1):
        add_para(doc, f"FR{i}: {f}")

    add_heading(doc, "4.2.2.2 Non Functional Requirements", 4)
    nfunc = [
        "Reproducible pipeline: notebooks run in numbered order with documented inputs and outputs.",
        "Security: API keys stored outside notebooks in .env and excluded from version control.",
        "Reliability: batch retries and checkpoint saves to handle network and API failures.",
        "Usability: clear section headers and charts readable by examiners without code knowledge.",
        "Maintainability: modular notebooks rather than one monolithic script.",
        "Transparency: limitations documented in the visualisation notebook.",
    ]
    for i, f in enumerate(nfunc, 1):
        add_para(doc, f"NFR{i}: {f}")

    add_heading(doc, "4.3 System Design", 2)
    add_heading(doc, "4.3.1 System Architecture", 3)
    add_para(
        doc,
        "System architecture describes how components connect to achieve the project goal. EF-02 uses a linear pipeline "
        "architecture where each stage reads processed output from the previous stage. External economic data joins at "
        "the consolidation stage.",
    )
    add_para(
        doc,
        "Pipeline flow: (1) 01_scraper collects raw headlines from Daily News and The Citizen into separate CSV files. "
        "(2) 02_cleaning merges, deduplicates, and normalises dates into tz_headlines_clean.csv. "
        "(3) 03_benchmarking compares TextBlob, VADER, FinBERT, and LLMs on a labelled sample. "
        "(4) 04_classifier performs single pass relevance, category, and sentiment labelling via Groq. "
        "(5) 05_consolidation aggregates to monthly level and merges CPI and FX data into Visualization_Data.csv. "
        "(6) 06_visualisation produces charts and Pearson correlation analysis.",
    )
    add_para(
        doc,
        "Key design decisions: models were benchmarked before full scraping to confirm LLM suitability; sentiment was "
        "classified by implied economic outcome rather than emotional tone; IPP Media was excluded due to archive and "
        "scraper constraints; API keys were moved to .env before any further development.",
    )
    doc.add_page_break()


def build_chapter5(doc: Document, fig_paths: dict[str, Path]) -> None:
    add_heading(doc, "CHAPTER FIVE: IMPLEMENTATION AND TESTING", 1)

    add_heading(doc, "5.1 Introduction", 2)
    add_para(
        doc,
        "This chapter describes how each part of the pipeline was implemented, presents key results from the classified "
        "dataset and correlation analysis, and reports testing outcomes for both model benchmarking and functional requirements.",
    )

    add_heading(doc, "5.2 Implementation", 2)

    add_heading(doc, "5.2.1 Web Scraper", 3)
    add_para(
        doc,
        "The scraper (01_scraper.ipynb) paginates through business category archives of Daily News and The Citizen. "
        "Each source has a test cell (3 to 5 pages) and a full scrape cell. Configuration for selectors and URL patterns "
        "is stored per source. The scraper saves checkpoint files every 30 pages, logs progress inline, and writes "
        "dailynews_raw.csv and citizen_raw.csv with columns: date, headline, url, source, scraped_on.",
    )

    add_heading(doc, "5.2.2 Data Cleaning", 3)
    add_para(
        doc,
        "The cleaning notebook (02_cleaning.ipynb) merges both raw CSV files, parses dates, removes duplicates and empty "
        "rows, and outputs tz_headlines_clean.csv with 5,310 unique headlines ready for classification.",
    )

    add_heading(doc, "5.2.3 Model Benchmarking", 3)
    add_para(
        doc,
        "The benchmarking notebook (03_benchmarking.ipynb) sent numbered batches of headlines to six models and compared "
        "output to human labelled sentiment. Results confirmed that LLMs far outperform rule based and FinBERT approaches "
        "on Tanzanian financial text.",
    )
    add_table(
        doc,
        ["Model", "Accuracy (%)"],
        [
            ["TextBlob", "28.2"],
            ["VADER", "50.0"],
            ["FinBERT", "66.7"],
            ["GPT-4o-mini", "92.9"],
            ["Llama 3.1 8B (Groq)", "92.7"],
            ["Llama 3.3 70B (Groq)", "93.5"],
        ],
        caption="Table 5.1: Model Benchmarking Results",
    )

    add_heading(doc, "5.2.4 Production Classifier", 3)
    add_para(
        doc,
        "The production classifier (04_classifier.ipynb) uses the Groq API with Llama models. Each API call processes "
        "25 headlines and returns a JSON array with relevance, category (one of nine: Forex, Policy, Banking, Trade, "
        "Agriculture, Energy, Transport, Investment, General), and sentiment (Positive, Negative, Neutral) based on "
        "implied economic outcome. Batch size is 25 with up to 3 retries. Temperature is set to 0 for consistent output. "
        "Of 5,310 headlines, 72.2% were marked relevant (3,836). Among relevant headlines, Neutral was most common (2,398), "
        "followed by Positive (1,114) and Negative (323). Top categories included Banking (785), Agriculture (593), "
        "Trade (580), General (551), and Policy (508).",
    )

    add_heading(doc, "5.2.5 Data Consolidation", 3)
    add_para(
        doc,
        "The consolidation notebook (05_consolidation.ipynb) drops irrelevant headlines, groups by YearMonth, and computes "
        "num_headlines, category shares, sentiment percentages, and dominant sentiment per month. It merges with monthly "
        "CPI inflation and average USD/TZS rate. The final dataset has 33 monthly rows from March 2022 to December 2024. "
        "A Net_Sentiment feature (Positive % minus Negative %) was engineered for analysis.",
    )

    add_heading(doc, "5.2.6 Visualisation and Analysis", 3)
    add_para(
        doc,
        "The visualisation notebook (06_visualisation.ipynb) presents eight sections: raw data summary, monthly overview, "
        "sentiment distribution, economic indicators, dual axis sentiment versus economy charts, Pearson correlation "
        "heatmap, a ranked correlation table with p values, and a limitations section. Inflation ranged from 2.96% to 4.94%. "
        "USD/TZS ranged from about 2,313 to 2,720 TZS per dollar. Mean relevant headlines per month was about 114.",
    )

    add_figure(doc, fig_paths["relevance"], "Figure 5.1: Headline relevance split (Relevant 72.2%, Irrelevant 27.8%)")
    add_figure(doc, fig_paths["sentiment"], "Figure 5.2: Monthly sentiment distribution (Neutral, Positive, Negative shares)")
    add_figure(doc, fig_paths["dual_axis"], "Figure 5.3: Positive and Negative sentiment versus USD/TZS exchange rate")
    add_figure(doc, fig_paths["correlation"], "Figure 5.4: Pearson correlation matrix for sentiment and economic variables")

    add_heading(doc, "5.3 Testing", 2)
    add_para(
        doc,
        "Testing verified that each pipeline stage worked as intended and that statistical outputs were correctly computed. "
        "Functional testing checks requirements against expected outcomes. Benchmarking testing compares model accuracy. "
        "Statistical validation uses Pearson correlation with p values.",
    )

    add_heading(doc, "5.3.1 Model Benchmarking Testing", 3)
    add_para(
        doc,
        "Six models were tested on the same labelled headline sample. Accuracy was measured as the share of predictions "
        "matching the ground truth sentiment label. LLMs exceeded 92% accuracy while TextBlob and VADER fell below 50%, "
        "confirming the decision to use Groq hosted LLMs for production classification.",
    )

    add_heading(doc, "5.3.2 Functional Testing", 3)
    add_table(
        doc,
        ["S/N", "Functional Requirement", "Expected Outcome", "Test Result"],
        [
            ["1", "Scraper saves headlines with date and source", "CSV rows with valid dates and source field", "Pass"],
            ["2", "Classifier returns valid JSON per batch", "25 labelled rows per successful batch", "Pass"],
            ["3", "Irrelevant headlines filtered from monthly data", "Only relevant=True used in aggregation", "Pass"],
            ["4", "Monthly merge aligns CPI and FX", "33 rows with all required columns", "Pass"],
            ["5", "Correlation statistics computed", "r and p printed for each sentiment pair", "Pass"],
            ["6", "API keys secured", "Keys loaded from .env, not hardcoded", "Pass"],
        ],
        caption="Table 5.2: Functional Testing Results",
    )

    add_heading(doc, "5.3.3 Statistical Validation", 3)
    add_para(
        doc,
        "Pearson correlation measured linear association between monthly sentiment variables and economic indicators. "
        "The strongest results involved the USD/TZS rate level:",
    )
    add_table(
        doc,
        ["Rank", "Pair", "Pearson r", "p value", "Significance"],
        [
            ["1", "Net Sentiment vs USD/TZS Rate", "0.583", "0.0004", "Highly significant (p ≤ 0.01)"],
            ["2", "Negative % vs USD/TZS Rate", "-0.484", "0.0043", "Highly significant (p ≤ 0.01)"],
            ["3", "Positive % vs USD/TZS Rate", "0.438", "0.0108", "Significant (p ≤ 0.05)"],
            ["4", "Net Sentiment vs Inflation %", "-0.413", "0.0170", "Significant (p ≤ 0.05)"],
            ["5", "Positive % vs Inflation %", "-0.328", "0.0622", "Not significant"],
            ["6", "Negative % vs Inflation %", "0.321", "0.0683", "Not significant"],
        ],
        caption="Table 5.3: Pearson Correlation Results (Sentiment vs Economic Indicators)",
    )
    add_para(
        doc,
        "All pairs involving month on month Rate_Change_% had p values above 0.05 and are not statistically significant "
        "at the 5% level. This suggests media sentiment tracks the level of the exchange rate more clearly than short "
        "term FX volatility.",
    )
    doc.add_page_break()


def build_chapter6(doc: Document) -> None:
    add_heading(doc, "CHAPTER SIX: CONCLUSION AND RECOMMENDATION", 1)

    add_heading(doc, "6.1 Introduction", 2)
    add_para(
        doc,
        "This chapter summarises what was achieved against each objective, discusses challenges and limitations, "
        "and offers recommendations for future work.",
    )

    add_heading(doc, "6.2 Conclusion", 2)

    add_heading(doc, "6.2.1 Classification Objective", 3)
    add_para(
        doc,
        "The project successfully classified 5,310 Tanzanian financial headlines using a Groq hosted large language model. "
        "Benchmarking showed LLM accuracy near 93%, far above rule based tools. The production run labelled 72.2% of "
        "headlines as economically relevant and assigned category and sentiment in a single API pass.",
    )

    add_heading(doc, "6.2.2 Monthly Aggregation Objective", 3)
    add_para(
        doc,
        "Monthly sentiment scores were computed for 33 months from March 2022 to December 2024. The merged dataset "
        "combines headline counts, category dominance, sentiment shares, inflation, and exchange rate in one table "
        "suitable for time series analysis.",
    )

    add_heading(doc, "6.2.3 Correlation Objective", 3)
    add_para(
        doc,
        "Statistical analysis found significant correlation between sentiment and the USD/TZS exchange rate. Net Sentiment "
        "correlated at r = 0.583 (p = 0.0004). Negative sentiment correlated at r = 0.484 (p = 0.0043) and Positive at "
        "r = 0.438 (p = 0.0108). Net Sentiment also correlated inversely with inflation at r = 0.413 (p = 0.017). These "
        "results support the research question that media tone relates to macro indicators, especially the exchange rate.",
    )

    add_heading(doc, "6.2.4 Visualisation Objective", 3)
    add_para(
        doc,
        "Charts showed neutral dominance in most months, a rising USD/TZS trend from 2023 onward alongside shifting "
        "sentiment shares, and visual alignment between positive sentiment and exchange rate movement in dual axis plots. "
        "The correlation heatmap confirmed that exchange rate level had the strongest associations.",
    )

    add_heading(doc, "6.2.5 Model Evaluation Objective", 3)
    add_para(
        doc,
        "LLMs proved suitable for Tanzanian financial sentiment when prompted to judge implied economic outcome. "
        "FinBERT and rule based models were insufficient. A later prompt tuning effort (v4) improved category accuracy "
        "from 59.6% to 79.9% on a 200 headline sample when using the smaller Llama 3.1 8B model, showing that prompt "
        "design matters as much as model choice.",
    )

    add_heading(doc, "6.3 Challenges and Limitations", 2)
    add_para(
        doc,
        "Several challenges affected the project. January and February 2022 were missing from the exchange rate source, "
        "reducing monthly observations from 36 to 33. Neutral sentiment dominated most months (about 62% of relevant "
        "headlines), which compressed the signal from positive and negative labels. The General category was over assigned "
        "in early classifier runs, though prompt updates were developed to address this. The production run used "
        "Llama 3.1 8B instead of the benchmarked 70B model due to token limits on the free API tier. Finally, correlation "
        "shows association, not causation, and 33 monthly points limit statistical power.",
    )

    add_heading(doc, "6.4 Recommendations", 2)
    add_para(
        doc,
        "Future work should re run full classification with the improved v5 prompt and expanded 11 category schema "
        "(adding Markets, Tourism, and Inflation while removing General). Additional news sources could be included using "
        "a headless browser for sites like IPP Media. The pipeline should be extended to BOT policy rate and GDP series. "
        "A simple dashboard could publish a monthly sentiment index for researchers. Publishing a labelled Tanzanian "
        "financial headline dataset would strengthen African NLP resources identified as lacking by Adelani et al. (2021).",
    )
    doc.add_page_break()

    add_heading(doc, "REFERENCES", 1)
    refs = [
        "Adelani, D. I., Abbott, J., Neubig, G., Dossou, B. F. P., Alabi, J. O., Masinde, W., ... & Buzaaba, H. (2021). MasakhaNER: Named entity recognition for African languages. Transactions of the Association for Computational Linguistics, 9, 1116–1131.",
        "Araci, D. (2019). FinBERT: Financial sentiment analysis with pre-trained language models. arXiv preprint arXiv:1908.10063.",
        "Baker, S. R., Bloom, N., & Davis, S. J. (2016). Measuring economic policy uncertainty. The Quarterly Journal of Economics, 131(4), 1593–1636.",
        "Bollen, J., Mao, H., & Zeng, X. (2011). Twitter mood predicts the stock market. Journal of Computational Science, 2(1), 1–8.",
        "Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2018). BERT: Pre-training of deep bidirectional transformers for language understanding. arXiv preprint arXiv:1810.04805.",
        "Hutto, C. J., & Gilbert, E. (2014). VADER: A parsimonious rule-based model for sentiment analysis of social media text. Proceedings of the International AAAI Conference on Web and Social Media, 8(1), 216–225.",
        "Loughran, T., & McDonald, B. (2011). When is a liability not a liability? Textual analysis, dictionaries, and 10-Ks. The Journal of Finance, 66(1), 35–65.",
        "Muriungi, P., & Kimani, D. (2020). Media sentiment and exchange rate volatility: Evidence from Kenyan financial news. African Journal of Economics and Finance, 8(2), 45–61.",
        "National Bureau of Statistics Tanzania. (2024). Consumer Price Index (CPI) summary data [Data set].",
        "Shapiro, A. H., Sudhof, M., & Wilson, D. J. (2022). Measuring news sentiment. Journal of Econometrics, 228(2), 221–243.",
        "Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market. The Journal of Finance, 62(3), 1139–1168.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(ref)
        run.font.name = FONT
        run.font.size = Pt(BODY_SIZE)

    doc.add_page_break()
    add_heading(doc, "APPENDIX", 1)
    add_para(
        doc,
        "No questionnaire was used in this project. Data collection relied on publicly available news archives and "
        "official economic statistics. Sample headline records are available in EF-02/data/processed/tz_headlines_labelled.csv.",
    )


def main() -> None:
    print("Exporting figures...")
    fig_paths = export_figures()

    print("Building document...")
    doc = Document()
    setup_styles(doc)

    build_cover(doc)
    build_front_matter(doc)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc, fig_paths)
    build_chapter6(doc)

    doc.save(str(OUTPUT))
    print(f"Report saved to: {OUTPUT}")
    print(f"File size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
