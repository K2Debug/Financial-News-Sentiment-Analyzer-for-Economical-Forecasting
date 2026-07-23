"""
Generate EF-02 Tool/Dashboard Project Report as a Word document.
Positions EF-02 as an upload → classify → analyse web tool (no scraping / notebooks).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
IMAGES = ROOT / "images"
OUTPUT = ROOT / "EF02_Tool_Report.docx"

FONT = "Times New Roman"
BODY_SIZE = 12


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size in [(1, 14), (2, 13), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        if level == 1:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_para(doc: Document, text: str, bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(BODY_SIZE)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None) -> None:
    if caption:
        cap = doc.add_paragraph()
        r = cap.add_run(caption)
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(BODY_SIZE)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = FONT
                run.font.size = Pt(BODY_SIZE)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.name = FONT
                    run.font.size = Pt(BODY_SIZE)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.8) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = FONT
    r.font.size = Pt(10)
    doc.add_paragraph()


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and select Update Field in Microsoft Word to refresh the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def build_cover(doc: Document) -> None:
    for text, bold in [
        ("ARDHI UNIVERSITY", True),
        ("DEPARTMENT OF COMPUTER SYSTEMS AND MATHEMATICS (CSM)", False),
        ("BACHELORS IN DATA SCIENCE AND ARTIFICIAL INTELLIGENCE (Bsc. DSAi)", False),
        (
            "DA 181: PROBLEM SOLVING USING FUNDAMENTAL CONCEPTS OF DATA SCIENCE AND ARTIFICIAL INTELLIGENCE.",
            False,
        ),
        ("", False),
        ("PROJECT TITLE: EF-02 FINANCIAL SENTIMENT ANALYSER DASHBOARD", True),
        ("", False),
        ("GROUP (5) MEMBERS:", True),
    ]:
        if text:
            add_para(doc, text, bold=bold, center=True)
        else:
            doc.add_paragraph()

    add_table(
        doc,
        ["S/N", "FULL NAMES", "REG.NO.", "SIGNATURE"],
        [
            ["1", "KITULA, KITULA-MASAGANYA ABDUL", "35552/T.2025", ""],
            ["2", "ILEKA, REVOCATUS OREST", "37058/T.2025", ""],
            ["3", "LYELLU, SABRINA RASHID", "37463/T.2025", ""],
            ["4", "IBRAHIM, GHADHAL NKRUMAH", "37254/T.2025", ""],
        ],
    )
    doc.add_page_break()


def build_front_matter(doc: Document) -> None:
    add_heading(doc, "DECLARATION", 1)
    add_para(
        doc,
        "We, the members of Group 5, hereby declare that this report is our own work and effort. "
        "The work in this report was carried out in accordance with the Regulations of Ardhi University. "
        "We have faithfully acknowledged all sources of information, and the tool described herein was "
        "designed and implemented by the group.",
    )
    add_table(
        doc,
        ["S/no", "Student's Name", "Signature", "Date"],
        [
            ["1", "KITULA, Kitula-Masaganya Abdul", "………………", "………….."],
            ["2", "ILEKA, Revocatus Orest", "………………", "………….."],
            ["3", "LYELLU, Sabrina Rashid", "………………", "………….."],
            ["4", "IBRAHIM, Ghadhal Nkrumah", "………………", "………….."],
        ],
    )
    doc.add_page_break()

    add_heading(doc, "CERTIFICATION", 1)
    add_para(
        doc,
        "The undersigned certify that they have read and hereby recommend for acceptance by Ardhi University "
        "a project titled: EF-02: Financial Sentiment Analyser Dashboard, in fulfillment of the requirements "
        "for the Bachelor of Science in Data Science and Artificial Intelligence.",
    )
    add_para(doc, "……………………")
    add_para(doc, "Dr. Winfrida")
    add_para(doc, "(Supervisor)")
    add_para(doc, "Date: ………….")
    doc.add_paragraph()
    add_para(doc, "……………………")
    add_para(doc, "Mr. Benitho")
    add_para(doc, "(Supervisor)")
    add_para(doc, "Date: ………….")
    doc.add_page_break()

    add_heading(doc, "ACKNOWLEDGMENT", 1)
    add_para(
        doc,
        "First and foremost, we thank Almighty God for giving us the strength and guidance to complete this "
        "project. We express our sincere gratitude to our supervisors, Mr. Benitho and Dr. Winfrida, for their "
        "guidance, constructive feedback, and encouragement throughout the development of the EF-02 dashboard.",
    )
    add_para(
        doc,
        "We also thank Ardhi University and the Department of Computer Systems and Mathematics for providing "
        "the learning environment and resources that made this work possible. We acknowledge publicly available "
        "economic statistics from the National Bureau of Statistics Tanzania and exchange-rate series used as "
        "example macro inputs in the tool demonstrations.",
    )
    doc.add_page_break()

    add_heading(doc, "LIST OF ABBREVIATIONS AND ACRONYMS", 1)
    add_table(
        doc,
        ["Abbreviation", "Meaning"],
        [
            ["API", "Application Programming Interface"],
            ["ARU", "Ardhi University"],
            ["BOT", "Bank of Tanzania"],
            ["CPI", "Consumer Price Index"],
            ["CSV", "Comma-Separated Values"],
            ["FX", "Foreign Exchange (USD/TZS)"],
            ["JSON", "JavaScript Object Notation"],
            ["LLM", "Large Language Model"],
            ["NBS", "National Bureau of Statistics"],
            ["NLP", "Natural Language Processing"],
            ["UI", "User Interface"],
            ["USD/TZS", "United States Dollar to Tanzanian Shilling"],
        ],
    )
    doc.add_page_break()

    add_heading(doc, "ABSTRACT", 1)
    add_para(
        doc,
        "This project delivers EF-02, a web-based Financial Sentiment Analyser dashboard for Tanzanian "
        "financial headlines. The tool lets users upload headline and macroeconomic CSV files, classify "
        "headlines with a large language model for relevance, topic category, and economic-outcome sentiment, "
        "then consolidate monthly aggregates and review interactive analysis views in a saved session.",
    )
    add_para(
        doc,
        "The implemented workflow follows a clear product schema: Data Import → Classify → Consolidate / "
        "Analyse. Classification uses OpenAI gpt-4o-mini through a FastAPI-backed dashboard with progress "
        "tracking, downloadable labelled outputs, and session history. Macro inputs (CPI inflation and USD/TZS "
        "rates) are provided by the user as separate uploads rather than being collected by the system.",
    )
    add_para(
        doc,
        "An example end-to-end GPT run on the dashboard demonstrates the tool in use: uploaded headline and "
        "macro CSVs covering an overlapping window from 2022 through 2024, batch classification with "
        "gpt-4o-mini, monthly consolidation, and analysis pages that show sentiment balances, net sentiment, "
        "inflation and exchange-rate series, and comparative charts. Model selection was guided by an earlier "
        "benchmark which showed that LLM labelling substantially outperformed TextBlob, VADER, and FinBERT "
        "on Tanzanian financial phrasing.",
    )
    add_para(
        doc,
        "The contribution of this project is a practical, reusable analyser: students and analysts can load "
        "their own headline files, obtain structured labels, merge them with macro series, and inspect results "
        "in a single session-based interface.",
    )
    doc.add_page_break()

    add_heading(doc, "TABLE OF CONTENTS", 1)
    add_toc_field(doc)
    doc.add_page_break()

    add_heading(doc, "LIST OF TABLES", 1)
    for t in [
        "Table 1.1: Value Proposition Canvas for EF-02",
        "Table 3.1: Methodology Table",
        "Table 4.1: Headline and Macro CSV Schemas",
        "Table 5.1: Model Benchmarking Results (classifier selection)",
        "Table 5.2: Functional Testing Results",
    ]:
        add_para(doc, t)
    doc.add_page_break()

    add_heading(doc, "LIST OF FIGURES", 1)
    for f in [
        "Figure 1.1: Value Proposition Canvas for EF-02",
        "Figure 4.1: EF-02 dashboard system architecture",
        "Figure 5.1: Dashboard home and saved sessions",
        "Figure 5.2: Data Import — uploading headlines, CPI, and USD/TZS",
        "Figure 5.3: Classify headlines with gpt-4o-mini",
        "Figure 5.4: Analysis summary panel after consolidation",
        "Figure 5.5: Monthly sentiment and macro time-series charts",
        "Figure 5.6: Sentiment versus inflation and USD/TZS comparative charts",
    ]:
        add_para(doc, f)
    doc.add_page_break()


def build_chapter_one(doc: Document) -> None:
    add_heading(doc, "CHAPTER ONE: INTRODUCTION", 1)
    add_heading(doc, "1.1 General Introduction", 2)
    add_para(
        doc,
        "Sentiment analysis is a Natural Language Processing (NLP) technique that assigns directional labels "
        "to text. In finance, sentiment tools help users summarise how coverage describes growth, risk, "
        "currency movement, or policy change. For Tanzania, English-language financial headlines discuss "
        "banking, trade, monetary policy, agriculture, energy, and investment — topics that matter to students, "
        "analysts, and decision makers who want a fast overview of media tone alongside official indicators.",
    )
    add_para(
        doc,
        "EF-02 is built as a software tool: a web dashboard that accepts CSV uploads, runs an LLM classifier, "
        "consolidates labelled headlines with CPI and USD/TZS series, and presents results in interactive charts. "
        "The product goal is repeatable analysis in a browser session, not a one-off research experiment.",
    )

    add_heading(doc, "1.2 Statement of the Problem", 2)
    add_para(
        doc,
        "Financial headlines are plentiful, but turning them into structured labels and linking them to "
        "macroeconomic series still usually requires scripting, scattered files, and manual charting. "
        "General-purpose lexicons and Western-trained finance models often misread East African economic "
        "phrasing, while calling APIs from disconnected scripts is hard to hand to a non-technical user.",
    )
    add_para(
        doc,
        "What is missing is a coherent local tool: upload headline and macro files, classify with an "
        "LLM tuned for economic-outcome sentiment, consolidate by month, and inspect outputs in one "
        "dashboard with downloadable artefacts and saved sessions.",
    )

    add_heading(doc, "1.3 Objectives", 2)
    add_heading(doc, "1.3.1 General Objective", 3)
    add_para(
        doc,
        "To design and implement a Financial Sentiment Analyser dashboard that classifies Tanzanian "
        "financial headlines and presents monthly sentiment views together with user-supplied CPI and "
        "USD/TZS data.",
    )

    add_heading(doc, "1.3.2 Specific Objectives", 3)
    for obj in [
        "To provide Data Import for headline CSVs (and CPI / USD/TZS series) with schema validation and session storage.",
        "To classify uploaded headlines for relevance, topic category, and economic-outcome sentiment using an LLM (gpt-4o-mini in the demonstrated build).",
        "To consolidate relevant labels into monthly aggregates and merge them with inflation and exchange-rate series.",
        "To present analysis outputs (summary panels and charts) in a browser UI with downloadable labelled and visualisation CSVs.",
        "To evaluate classifier approach suitability (rule-based, FinBERT, LLM) as an engineering choice for the labelling component.",
    ]:
        add_para(doc, obj)

    add_heading(doc, "1.3.3 Value Proposition Canvas", 3)
    add_para(
        doc,
        "A Value Proposition Canvas links user needs to what the product offers. For EF-02, users need a "
        "fast way to label financial headlines and view them beside macroeconomic context without building "
        "a pipeline from scratch. The dashboard answers that need with upload, classification, consolidation, "
        "and analysis in one session-based tool.",
    )
    add_para(doc, "Figure 1.1: Value proposition canvas for EF-02", center=True)
    add_table(
        doc,
        ["Canvas side", "Items"],
        [
            [
                "Customer jobs",
                "Label financial headlines; track monthly tone; overlay inflation/FX; reuse past runs",
            ],
            [
                "Pains",
                "Manual labelling; fragile scripts; Western lexicons; no single UI for Tanzanian text + macros",
            ],
            [
                "Gains",
                "Batch LLM labels; 11-category schema; cached sessions; downloads; charts in one place",
            ],
            [
                "Products & services",
                "EF-02 Sentiment Dashboard (Import → Classify → Analyse)",
            ],
            [
                "Pain relievers",
                "CSV upload workflow, progress UI, settings for model/API, optional pre-labelled skip path",
            ],
            [
                "Gain creators",
                "Structured JSON labels, monthly consolidation, interactive analysis screens",
            ],
        ],
        caption="Table 1.1: Value Proposition Canvas for EF-02",
    )

    add_heading(doc, "1.4 System Capability Questions", 2)
    for q in [
        "CQ1: Can the dashboard accept validated headline and macro CSV uploads and store them in a named session?",
        "CQ2: Can an LLM classifier assign relevance, category, and economic sentiment to Tanzanian financial headlines in batch with progress and downloadable labels?",
        "CQ3: Can the system consolidate monthly sentiment shares and merge them with CPI and USD/TZS series?",
        "CQ4: Does the Analysis page present usable summary metrics and charts for an end-to-end GPT run?",
        "CQ5: Is an LLM-based classifier a stronger engineering choice than rule-based or FinBERT baselines for this labelling task?",
    ]:
        add_para(doc, q)

    add_heading(doc, "1.5 Significance of the Study", 2)
    add_para(doc, "Upon completing this project, the following benefits are expected:")
    for b in [
        "It delivers a reusable software tool for Tanzanian financial headline sentiment analysis.",
        "It demonstrates an end-user workflow (upload → classify → analyse) suitable for students and analysts.",
        "It shows how prompt-configured LLMs can handle local economic phrasing better than lexicon or Western FinBERT defaults.",
        "It packages labelling, consolidation, and visualisation so results are inspectable without writing custom scripts for each run.",
        "It builds practical skills in API integration, web application design, data validation, and interactive visualisation.",
    ]:
        add_para(doc, b)

    add_heading(doc, "1.6 Structure of the Report", 2)
    for s in [
        "Chapter One introduces the problem, tool objectives, capability questions, and significance.",
        "Chapter Two reviews related literature and the product gap that motivates a dashboard tool.",
        "Chapter Three explains the methodology used to build and validate the system.",
        "Chapter Four presents requirements and system design for the dashboard.",
        "Chapter Five describes implementation, the example GPT run, and testing.",
        "Chapter Six provides conclusions, limitations, and product recommendations.",
    ]:
        add_para(doc, s)
    doc.add_page_break()


def build_chapter_two(doc: Document) -> None:
    add_heading(doc, "CHAPTER TWO: LITERATURE REVIEW", 1)
    add_heading(doc, "2.1 Introduction", 2)
    add_para(
        doc,
        "This chapter reviews selected work on financial text sentiment and language models. The goal is not "
        "to reproduce prior correlation studies, but to justify why EF-02 uses an LLM-based classifier inside "
        "a practical upload-and-analyse tool for Tanzanian financial headlines.",
    )

    add_heading(doc, "2.2 Related Studies", 2)
    sections = [
        (
            "2.2.1 Finance-Specific Sentiment Lexicon (Loughran and McDonald, 2011)",
            "Loughran and McDonald (2011) showed that general dictionaries mislabel financial English because "
            "words such as liability or risk carry domain-specific meaning. Lexicon tools therefore need "
            "domain care — a reason EF-02 does not rely on generic polarity lists for production labelling.",
        ),
        (
            "2.2.2 Media Tone and Markets (Tetlock, 2007)",
            "Tetlock (2007) linked press tone to market behaviour, illustrating why media text is used as an "
            "economic signal. EF-02 treats headlines as inputs to a labelling tool that users can combine with "
            "their own macro files.",
        ),
        (
            "2.2.3 Public Mood and Markets (Bollen, Mao and Zeng, 2011)",
            "Bollen, Mao, and Zeng (2011) used social-media mood features for market prediction. Their work "
            "supports computational sentiment as a product class, while EF-02 focuses on financial news "
            "headlines and a dashboard workflow.",
        ),
        (
            "2.2.4 BERT (Devlin et al., 2018)",
            "BERT introduced bidirectional transformers that improve contextual understanding over bag-of-words "
            "models and underpin later domain models such as FinBERT.",
        ),
        (
            "2.2.5 FinBERT (Araci, 2019)",
            "FinBERT fine-tunes BERT on Western financial text. It improves on generic NLP for earnings-style "
            "English, but still underperformed LLM prompting on our Tanzanian headline sample — informing "
            "EF-02’s choice of an LLM classifier for the tool’s labelling stage.",
        ),
        (
            "2.2.6 News Tone Indices (Shapiro, Sudhof, and Wilson, 2022)",
            "Shapiro et al. (2022) aggregate newspaper tone into macroeconomic indices. EF-02 similarly "
            "aggregates labelled headlines by month, but as a configurable dashboard feature for user data.",
        ),
        (
            "2.2.7 East African Financial News (Muriungi and Kimani, 2020)",
            "Regional work on Kenyan financial media sentiment highlights the value of African-context text. "
            "EF-02 targets Tanzanian English financial headlines with category and economic-outcome labels.",
        ),
    ]
    for title, body in sections:
        add_heading(doc, title, 3)
        add_para(doc, body)

    add_heading(doc, "2.3 Product Gap", 2)
    add_para(
        doc,
        "Prior work establishes that financial sentiment is useful and that domain or LLM methods beat naive "
        "lexicons. The remaining gap for this course project is productisation: an end-user dashboard where "
        "Tanzanian headline CSVs can be uploaded, classified, consolidated with CPI and FX files, and reviewed "
        "visually in one place. EF-02 addresses that gap."
    )
    doc.add_page_break()


def build_chapter_three(doc: Document) -> None:
    add_heading(doc, "CHAPTER THREE: METHODOLOGY", 1)
    add_heading(doc, "3.1 Introduction", 2)
    add_para(
        doc,
        "This chapter describes how EF-02 was planned, designed, implemented, and tested as a software tool. "
        "Methods follow a practical systems approach suitable for a dashboard product with an NLP labelling core.",
    )

    add_heading(doc, "3.2 Methodological Steps", 2)
    add_heading(doc, "3.2.1 General Methodology", 3)
    add_para(
        doc,
        "Development followed an iterative build-and-test cycle: define the upload → classify → analyse "
        "workflow, implement each stage behind a FastAPI service with a browser UI, validate CSV schemas, "
        "run classification against the OpenAI API, consolidate monthly outputs, and verify functional behaviour "
        "on an end-to-end GPT demonstration run.",
    )

    add_heading(doc, "3.2.2 Gathering User Requirements", 3)
    add_para(
        doc,
        "Requirements were elicited by reviewing how analysts would use the tool: provide headline files, "
        "attach official CPI and FX CSVs, obtain labelled downloads, and view charts. Supervisory feedback "
        "emphasised a clear product schema and a dashboard presentation suited to end users.",
    )

    add_heading(doc, "3.2.3 System Design Methodology", 3)
    add_para(
        doc,
        "Design used modular separation: presentation (HTML/CSS/JS dashboard), application services "
        "(job/session orchestration), and core processing (validation, classification, consolidation, "
        "analysis helpers). Sessions persist artefacts under a runs directory so users can reopen work.",
    )

    add_heading(doc, "3.2.4 Implementation Methodology", 3)
    add_para(
        doc,
        "Implementation used Python (FastAPI, pandas, OpenAI SDK) for the backend and a static SPA-style "
        "frontend for Import, Classify, and Analysis pages. Classification prompts request structured JSON "
        "with relevance, one of eleven topic categories, and Positive/Negative/Neutral economic sentiment.",
    )

    add_heading(doc, "3.2.5 Testing and Validation", 3)
    add_para(
        doc,
        "Validation combined functional tests of each dashboard stage, schema checks on uploads, and a "
        "classifier benchmark comparing lexicon, FinBERT, and LLM accuracy on a labelled sample to justify "
        "the production labelling choice.",
    )

    add_heading(doc, "3.3 Methodology Table", 2)
    add_table(
        doc,
        ["Specific Objective", "Methodology", "Tools", "Deliverable"],
        [
            [
                "Import headline and macro data",
                "CSV upload with column validation and session storage",
                "FastAPI, pandas",
                "Session files: headlines, cpi.csv, fx.csv",
            ],
            [
                "Classify headlines",
                "Batch LLM classification with structured JSON schema",
                "OpenAI gpt-4o-mini, dashboard progress UI",
                "tz_headlines_labelled.csv",
            ],
            [
                "Consolidate monthly views",
                "Aggregate relevant labels by YearMonth; merge CPI and FX",
                "pandas consolidation module",
                "Visualization_Data.csv",
            ],
            [
                "Present analysis",
                "Summary panels and interactive charts in Analysis page",
                "Chart.js, dashboard UI",
                "Analysis screens / downloads",
            ],
            [
                "Select labelling approach",
                "Benchmark lexicon, FinBERT, and LLM accuracy",
                "Labelled sample set, sklearn metrics",
                "Benchmark table; LLM adopted for tool",
            ],
            [
                "Verify system behaviour",
                "Functional testing of Import / Classify / Analyse",
                "Manual UI tests on GPT demo session",
                "Table 5.2 test results",
            ],
        ],
        caption="Table 3.1: Methodology Table",
    )
    doc.add_page_break()


def build_chapter_four(doc: Document) -> None:
    add_heading(doc, "CHAPTER FOUR: SYSTEM ANALYSIS AND DESIGN", 1)
    add_heading(doc, "4.1 Introduction", 2)
    add_para(
        doc,
        "This chapter specifies what the EF-02 dashboard must do and how its components are organised. "
        "Requirements assume user-supplied CSVs and an external LLM API for classification.",
    )

    add_heading(doc, "4.2 Requirement Analysis", 2)
    add_heading(doc, "4.2.1 System Requirements", 3)
    add_heading(doc, "4.2.1.1 Data Requirements", 4)
    add_para(
        doc,
        "The tool expects three primary uploads for the standard path, with an optional pre-labelled path "
        "to skip classification when labels already exist.",
    )
    add_table(
        doc,
        ["Dataset", "Required columns", "Role"],
        [
            ["Headlines", "date, headline", "Raw inputs for Classify"],
            [
                "Labelled CSV (optional)",
                "date, headline, relevant, category, sentiment",
                "Skip Classify; go to Analyse",
            ],
            ["CPI", "date, inflation_rate_pct", "Monthly inflation context"],
            ["USD/TZS", "Date, Price", "Exchange-rate context"],
        ],
        caption="Table 4.1: Headline and Macro CSV Schemas",
    )

    add_heading(doc, "4.2.1.2 Software Requirements", 4)
    add_para(
        doc,
        "Python 3 with FastAPI and Uvicorn; pandas; OpenAI Python SDK; a modern browser for the dashboard; "
        "and an OpenAI API key configured in Settings or environment variables.",
    )

    add_heading(doc, "4.2.1.3 Hardware Requirements", 4)
    add_para(
        doc,
        "A standard workstation or laptop able to run the local FastAPI server and open the dashboard in a "
        "browser. Network access is required for OpenAI classification calls.",
    )

    add_heading(doc, "4.2.2.1 Functional Requirements", 4)
    for fr in [
        "FR1: Create and manage named analysis sessions (open, rename, delete).",
        "FR2: Upload and validate headline, CPI, and USD/TZS CSV files (and optional labelled CSV).",
        "FR3: Run batch classification with progress, retries, and downloadable labelled CSV.",
        "FR4: Consolidate monthly aggregates and produce visualisation data.",
        "FR5: Display analysis summary metrics and charts on the Analysis page.",
        "FR6: Allow Settings for model selection, prompts, and API key.",
    ]:
        add_para(doc, fr)

    add_heading(doc, "4.2.2.2 Non-Functional Requirements", 4)
    for nfr in [
        "NFR1: Usability — clear Import → Classify → Analyse navigation with locked steps until prerequisites exist.",
        "NFR2: Reliability — checkpoints/progress messaging during classification; cached artefacts between stages.",
        "NFR3: Maintainability — core logic separated into reusable modules behind the API.",
        "NFR4: Portability — local server deployment with static frontend assets.",
    ]:
        add_para(doc, nfr)

    add_heading(doc, "4.3 System Design", 2)
    add_heading(doc, "4.3.1 System Architecture", 3)
    add_para(
        doc,
        "Figure 4.1 summarises the architecture. The browser UI talks to FastAPI job routes. Jobs store "
        "uploads and outputs under a per-session folder. Classification calls OpenAI; consolidation and "
        "analysis helpers run locally on labelled and macro CSVs.",
    )
    add_para(doc, "Figure 4.1: EF-02 dashboard system architecture", center=True)
    add_para(
        doc,
        "[Browser Dashboard UI] → [FastAPI: jobs, uploads, events, settings] → "
        "[ef02_core: validate / classify / consolidate / correlations] → [OpenAI gpt-4o-mini]. "
        "Session artefacts live under runs/{job_id}/.",
        center=True,
    )
    doc.add_page_break()


def build_chapter_five(doc: Document) -> None:
    add_heading(doc, "CHAPTER FIVE: IMPLEMENTATION AND TESTING", 1)
    add_heading(doc, "5.1 Introduction", 2)
    add_para(
        doc,
        "This chapter presents the implemented EF-02 Sentiment Dashboard and demonstrates an example "
        "gpt-4o-mini run using the screenshots in the accompanying images folder. Testing confirms that "
        "each stage of the upload → classify → analyse schema behaves as designed.",
    )

    add_heading(doc, "5.2 Implementation", 2)
    add_heading(doc, "5.2.1 Home and Session Management", 3)
    add_para(
        doc,
        "The home page introduces the three-step workflow and lists saved sessions with status, headline "
        "counts, and open/rename/delete actions. Users start a named analysis or resume a previous session.",
    )
    add_figure(doc, IMAGES / "home.png", "Figure 5.1: Dashboard home and saved sessions")

    add_heading(doc, "5.2.2 Data Import", 3)
    add_para(
        doc,
        "Data Import accepts headline CSVs (date, headline), CPI (date, inflation_rate_pct), and USD/TZS "
        "(Date, Price). Multiple headline files can be merged. After successful uploads, the UI shows the "
        "overlap window and headline count, then unlocks Classify. Users with a pre-labelled CSV can also "
        "skip classification and proceed to Analysis once macro files are present.",
    )
    add_figure(
        doc,
        IMAGES / "upload.png",
        "Figure 5.2: Data Import — uploading headlines, CPI, and USD/TZS",
    )

    add_heading(doc, "5.2.3 Classification", 3)
    add_para(
        doc,
        "The Classify page runs OpenAI gpt-4o-mini with model and API key configured under Settings. "
        "Headlines are processed in batches with an elapsed timer, progress bar, and activity log. "
        "Each successful run yields a labelled CSV with relevance, category, and sentiment fields that can "
        "be downloaded from the same screen.",
    )
    add_figure(doc, IMAGES / "classifying.png", "Figure 5.3: Classify headlines with gpt-4o-mini")

    add_heading(doc, "5.2.4 Consolidation and Analysis Views", 3)
    add_para(
        doc,
        "After labels exist, Analysis consolidates relevant headlines by month, merges CPI and FX series, "
        "and renders a summary panel plus interactive charts. The example GPT run shows a 36-month window "
        "with Neutral and Positive months alternating as dominant, comparatively low average negative share, "
        "and visual overlays of inflation and USD/TZS movement. Correlation outcome rows appear as part of "
        "the Analysis UI for user inspection; they are a feature of the tool’s display, not a research verdict "
        "claimed by this report.",
    )
    add_figure(
        doc,
        IMAGES / "report conclusion.png",
        "Figure 5.4: Analysis summary panel after consolidation (example GPT run)",
    )
    add_figure(
        doc,
        IMAGES / "report grapghs 1.png",
        "Figure 5.5: Monthly sentiment and macro time-series charts (example GPT run)",
        width=5.8,
    )
    add_figure(
        doc,
        IMAGES / "report grapghs 2.png",
        "Figure 5.6: Sentiment versus inflation and USD/TZS comparative charts (example GPT run)",
        width=5.8,
    )

    add_heading(doc, "5.2.5 Classifier Selection Note", 3)
    add_para(
        doc,
        "Before locking the labelling component, candidate approaches were scored on a labelled sample. "
        "Rule-based tools and FinBERT underperformed relative to LLM prompting on Tanzanian financial "
        "phrasing. The dashboard therefore uses an LLM classifier (gpt-4o-mini in the demonstrated build).",
    )
    add_table(
        doc,
        ["Model", "Approximate accuracy"],
        [
            ["TextBlob", "~28%"],
            ["VADER", "~50%"],
            ["FinBERT", "~67%"],
            ["LLM (prompted)", "~92–93%"],
        ],
        caption="Table 5.1: Model Benchmarking Results (classifier selection)",
    )

    add_heading(doc, "5.3 Testing", 2)
    add_heading(doc, "5.3.1 Functional Testing", 3)
    add_para(
        doc,
        "Functional testing verified each major dashboard capability on the example GPT session. Results "
        "are summarised in Table 5.2.",
    )
    add_table(
        doc,
        ["S/N", "Functional requirement", "Expected outcome", "Test result"],
        [
            ["1", "Session management", "Create/open/rename/delete sessions", "Passed"],
            ["2", "Data import", "Validate and store headline, CPI, FX CSVs", "Passed"],
            ["3", "Classification", "Produce labelled CSV via gpt-4o-mini batches", "Passed"],
            ["4", "Consolidation", "Build monthly merged visualisation dataset", "Passed"],
            ["5", "Analysis UI", "Show summary metrics and charts", "Passed"],
            ["6", "Downloads", "Export labelled and visualisation CSVs", "Passed"],
        ],
        caption="Table 5.2: Functional Testing Results",
    )
    add_para(
        doc,
        "Result: All functional requirements exercised in the demonstration run produced outputs consistent "
        "with the intended tool behaviour.",
    )

    add_heading(doc, "5.3.2 Usability Observation", 3)
    add_para(
        doc,
        "The locked navigation pattern (Import before Classify, labels and macros before Analysis) kept the "
        "schema clear during the demo. Progress feedback during classification made long batch runs observable "
        "without leaving the page.",
    )
    doc.add_page_break()


def build_chapter_six(doc: Document) -> None:
    add_heading(doc, "CHAPTER SIX: CONCLUSION AND RECOMMENDATION", 1)
    add_heading(doc, "6.1 Introduction", 2)
    add_para(
        doc,
        "This chapter evaluates EF-02 against its tool objectives, notes practical limitations of the "
        "dashboard build, and recommends product next steps.",
    )

    add_heading(doc, "6.2 Conclusion", 2)
    add_heading(doc, "6.2.1 Data Import Objective", 3)
    add_para(
        doc,
        "The dashboard accepts validated headline, CPI, and USD/TZS CSV uploads into named sessions. The "
        "example GPT run successfully loaded overlapping 2022–2024 inputs and proceeded through the schema.",
    )

    add_heading(doc, "6.2.2 Classification Objective", 3)
    add_para(
        doc,
        "Batch classification with gpt-4o-mini produces structured relevance, category, and economic-outcome "
        "sentiment labels, with progress UI and downloadable labelled CSV. Benchmark evidence supported "
        "choosing an LLM over lexicon and FinBERT baselines for this labelling component.",
    )

    add_heading(doc, "6.2.3 Consolidation Objective", 3)
    add_para(
        doc,
        "Relevant labels are aggregated by month and merged with inflation and exchange-rate series into a "
        "visualisation dataset ready for the Analysis page.",
    )

    add_heading(doc, "6.2.4 Analysis and Visualisation Objective", 3)
    add_para(
        doc,
        "The Analysis page delivers summary metrics and interactive charts (sentiment balances, net "
        "sentiment, inflation, USD/TZS, and comparative views) for the example GPT run, confirming that the "
        "tool presents consolidated outputs in a usable interface.",
    )

    add_heading(doc, "6.2.5 Overall Tool Delivery", 3)
    add_para(
        doc,
        "EF-02 meets its primary goal as a Financial Sentiment Analyser dashboard: upload → classify → "
        "consolidate / analyse, with session history and downloads. The project’s conclusion is that the "
        "tool works end-to-end for Tanzanian financial headline workflows — not that a particular sentiment–"
        "macro statistical relationship has been proven or rejected.",
    )

    add_heading(doc, "6.3 Challenges and Limitations", 2)
    add_para(
        doc,
        "Key limitations include dependence on an external LLM API (cost, rate limits, and connectivity); "
        "sensitivity of labels to prompt wording; need for clean CSV schemas from users; and the fact that "
        "dashboard charts reflect the quality of uploaded data and labels. The tool also requires a local "
        "server process and an API key for classification runs.",
    )

    add_heading(doc, "6.4 Recommendations", 2)
    add_para(
        doc,
        "Future product work should consider a distilled local classifier for deterministic offline labelling, "
        "stronger upload templates and validation messages, broader prompt presets for different sectors, and "
        "UI polish around long-running jobs. Supporting more flexible date formats and richer export packs "
        "(charts + CSV + summary) would further improve the analyser as a day-to-day tool.",
    )
    doc.add_page_break()


def build_references_appendix(doc: Document) -> None:
    add_heading(doc, "REFERENCES", 1)
    refs = [
        "Araci, D. (2019). FinBERT: Financial sentiment analysis with pre-trained language models. arXiv preprint arXiv:1908.10063.",
        "Bollen, J., Mao, H., & Zeng, X. (2011). Twitter mood predicts the stock market. Journal of Computational Science, 2(1), 1–8.",
        "Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2018). BERT: Pre-training of deep bidirectional transformers for language understanding. arXiv preprint arXiv:1810.04805.",
        "Hutto, C. J., & Gilbert, E. (2014). VADER: A parsimonious rule-based model for sentiment analysis of social media text. Proceedings of the International AAAI Conference on Web and Social Media.",
        "Loughran, T., & McDonald, B. (2011). When is a liability not a liability? Textual analysis, dictionaries, and 10-Ks. The Journal of Finance, 66(1), 35–65.",
        "Muriungi, P., & Kimani, D. (2020). Media sentiment and exchange rate volatility: Evidence from Kenyan financial news. African Journal of Economics and Finance, 8(2), 45–61.",
        "National Bureau of Statistics Tanzania. (2024). Consumer Price Index (CPI) summary data [Data set].",
        "Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner’s Approach (9th ed.). McGraw-Hill.",
        "Shapiro, A. H., Sudhof, M., & Wilson, D. J. (2022). Measuring news sentiment. Journal of Econometrics, 228(2), 221–243.",
        "Sommerville, I. (2016). Software Engineering (10th ed.). Pearson Education.",
        "Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market. The Journal of Finance, 62(3), 1139–1168.",
    ]
    for r in refs:
        add_para(doc, r)

    doc.add_page_break()
    add_heading(doc, "APPENDIX", 1)
    add_para(
        doc,
        "No questionnaire was used in this project. Demonstration data consisted of user-uploaded headline "
        "and macroeconomic CSV files processed through the EF-02 Sentiment Dashboard. Sample column schemas "
        "are listed in Table 4.1. Screenshots of the example gpt-4o-mini run are included in Chapter Five.",
    )


def main() -> None:
    doc = Document()
    setup_styles(doc)
    build_cover(doc)
    build_front_matter(doc)
    build_chapter_one(doc)
    build_chapter_two(doc)
    build_chapter_three(doc)
    build_chapter_four(doc)
    build_chapter_five(doc)
    build_chapter_six(doc)
    build_references_appendix(doc)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
